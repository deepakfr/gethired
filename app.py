import streamlit as st
import pdfplumber
import docx
from docx import Document
from docx.shared import Pt
from fpdf import FPDF
import requests
import smtplib
from email.message import EmailMessage
from io import BytesIO
import unicodedata

# ✅ Set your Groq API Key
GROQ_API_KEY = "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"
GROQ_ENDPOINT = "https://api.groq.com/openai/v1/chat/completions"

# ✅ Set your Email credentials
SMTP_SERVER = 'smtp.gmail.com'
SMTP_PORT = 465
SENDER_EMAIL = 'yourcompanyemail@gmail.com'  # <-- Your Gmail
SENDER_PASSWORD = 'yourapppassword'           # <-- Gmail App Password

# 🧠 Extract Text from uploaded Resume
def extract_text(file):
    if file.type == "application/pdf":
        with pdfplumber.open(file) as pdf:
            text = ""
            for page in pdf.pages:
                if page.extract_text():
                    text += page.extract_text() + "\n"
        return text
    elif file.type in ["application/vnd.openxmlformats-officedocument.wordprocessingml.document", "application/msword"]:
        doc = docx.Document(file)
        return "\n".join([p.text for p in doc.paragraphs])
    else:
        return file.read().decode("utf-8")

# 🧠 Unicode Safe
def sanitize_text(text):
    return unicodedata.normalize('NFKD', text).encode('latin-1', 'ignore').decode('latin-1')

# 🧠 Auto-extract Top Info (Name, Job Title, Contact Info)
def extract_top_info(resume_text):
    lines = resume_text.strip().split('\n')
    lines = [l.strip() for l in lines if l.strip()]
    
    name = lines[0] if len(lines) > 0 else ""
    job_title = lines[1] if len(lines) > 1 else ""
    contact_info = lines[2] if len(lines) > 2 else ""
    
    return name, job_title, contact_info

# 🧠 Call Groq API to tailor Resume & Cover Letter
def tailor_resume_and_coverletter(existing_resume, job_description):
    prompt = f"""
    Act as a professional resume writer and career expert.

    Based on the following current resume:
    {existing_resume}

    And the following job description:
    {job_description}

    Strict Instructions:
    - Rewrite the resume fully to match the job description.
    - Structure Resume into these sections:
      1. Profile Summary
      2. Languages
      3. Skills
      4. Expertise Areas
      5. Academic Projects
      6. Work Experience
      7. Education
      8. Soft Skills
    - Bullet points start with action verbs.
    - Clean ATS-optimized style, no personal pronouns.

    Then write a Cover Letter:
    - 3 paragraphs: Introduction, Skills match, Conclusion.

    Output format:
    ### Resume
    [resume here]

    ### Cover Letter
    [cover letter here]
    """
    headers = {
        "Authorization": f"Bearer {GROQ_API_KEY}",
        "Content-Type": "application/json"
    }
    payload = {
        "model": "llama3-8b-8192",
        "messages": [{"role": "system", "content": "You are a helpful AI assistant."},
                     {"role": "user", "content": prompt}],
        "temperature": 0.3
    }
    response = requests.post(GROQ_ENDPOINT, headers=headers, json=payload)
    if response.status_code == 200:
        result = response.json()
        return result["choices"][0]["message"]["content"]
    else:
        raise Exception(f"❌ Error from Groq API: {response.status_code} - {response.text}")

# 📄 Create DOCX with centered info
def create_docx(name, job_title, contact_info, resume_body, cover_letter):
    doc = Document()

    # Centered Name
    name_para = doc.add_paragraph()
    name_para.alignment = 1
    run = name_para.add_run(name.upper())
    run.bold = True
    run.font.size = Pt(18)

    # Centered Job Title
    title_para = doc.add_paragraph()
    title_para.alignment = 1
    run = title_para.add_run(job_title.title())
    run.font.size = Pt(14)

    # Centered Contact Info
    contact_para = doc.add_paragraph()
    contact_para.alignment = 1
    run = contact_para.add_run(contact_info)
    run.font.size = Pt(10)

    doc.add_paragraph("\n")  # Space

    # Resume Body
    for line in resume_body.split("\n"):
        if line.strip():
            if line.startswith("-"):
                doc.add_paragraph(line.strip(), style='ListBullet')
            else:
                doc.add_paragraph(line.strip())

    doc.add_page_break()

    # Cover Letter
    doc.add_heading('Cover Letter', 0)
    for line in cover_letter.split("\n"):
        if line.strip():
            doc.add_paragraph(line.strip())

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream

# 📄 Create Beautiful PDF
class BeautifulPDF(FPDF):
    def header(self):
        self.set_font('Helvetica', 'B', 20)
        self.set_text_color(30, 30, 30)
        self.cell(0, 15, 'GetHired - Resume & Cover Letter', ln=True, align='C')
        self.ln(10)

    def add_section_title(self, title):
        self.set_font('Helvetica', 'B', 16)
        self.set_text_color(0, 102, 204)
        self.cell(0, 10, title, ln=True)
        self.ln(5)

    def add_body_text(self, content):
        self.set_font('Helvetica', '', 12)
        self.set_text_color(50, 50, 50)
        self.multi_cell(0, 8, content)
        self.ln(8)

def create_beautiful_pdf(resume, cover_letter):
    pdf = BeautifulPDF()
    pdf.add_page()
    pdf.add_section_title("Tailored Resume")
    pdf.add_body_text(sanitize_text(resume))
    pdf.add_page()
    pdf.add_section_title("Cover Letter")
    pdf.add_body_text(sanitize_text(cover_letter))
    return pdf.output(dest='S').encode('latin1')

# 📧 Send Email
def send_email_with_attachments(to_email, docx_data, pdf_data):
    msg = EmailMessage()
    msg['Subject'] = "Your Tailored Resume & Cover Letter from GetHired 🎯"
    msg['From'] = SENDER_EMAIL
    msg['To'] = to_email

    msg.set_content("Hi there!\n\nPlease find attached your Tailored Resume and Cover Letter. Good luck! 🚀\n\n- GetHired Team")

    msg.add_attachment(docx_data.getvalue(), maintype='application', subtype='vnd.openxmlformats-officedocument.wordprocessingml.document', filename="Tailored_Resume_and_CoverLetter.docx")
    msg.add_attachment(pdf_data, maintype='application', subtype='pdf', filename="Tailored_Resume_and_CoverLetter.pdf")

    with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT) as smtp:
        smtp.login(SENDER_EMAIL, SENDER_PASSWORD)
        smtp.send_message(msg)

# 🏠 Streamlit App
st.set_page_config(page_title="GetHired - Tailor My Resume", page_icon="📝")
st.title("📝 GetHired - Tailor My Resume")
st.caption("Upload Resume ➔ Paste Job ➔ Auto-detect Info ➔ Tailored DOCX + Beautiful PDF ➔ Email to yourself!")

st.markdown("---")

# Upload and Input
st.subheader("📄 Upload Your Current Resume")
existing_resume_file = st.file_uploader("Upload Resume (PDF, DOCX, or TXT)", type=["pdf", "docx", "txt"])

if existing_resume_file:
    existing_resume = extract_text(existing_resume_file)
    name, job_title, contact_info = extract_top_info(existing_resume)
    
    st.success("✅ Info detected from Resume!")

    st.subheader("✏️ Personal Information (Auto-filled, editable)")
    name = st.text_input("Full Name", value=name)
    job_title = st.text_input("Job Title", value=job_title)
    contact_info = st.text_area("Contact Info", value=contact_info)

    st.markdown("---")
    st.subheader("🖊️ Paste Job Description")
    job_description_text = st.text_area("Paste the Job Description here...")

    if st.button("🚀 Tailor Resume & Create Documents"):
        if job_description_text.strip():
            with st.spinner("✍️ Tailoring your Resume and writing your Cover Letter..."):
                try:
                    output = tailor_resume_and_coverletter(existing_resume, job_description_text)

                    if "### Cover Letter" in output:
                        tailored_resume, cover_letter = output.split("### Cover Letter")
                        tailored_resume = tailored_resume.replace("### Resume", "").strip()
                        cover_letter = cover_letter.strip()

                        # Create files
                        docx_file = create_docx(name, job_title, contact_info, tailored_resume, cover_letter)
                        pdf_file = create_beautiful_pdf(tailored_resume, cover_letter)

                        st.success("✅ Documents ready!")

                        st.download_button("📥 Download DOCX", data=docx_file, file_name="Tailored_Resume_and_CoverLetter.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                        st.download_button("📥 Download Beautiful PDF", data=pdf_file, file_name="Tailored_Resume_and_CoverLetter.pdf", mime="application/pdf")

                        # 📩 Ask for Email
                        st.markdown("---")
                        st.subheader("📩 Email Your Documents")
                        user_email = st.text_input("Enter your Email address:")

                        if st.button("📤 Send to My Email"):
                            if user_email:
                                send_email_with_attachments(user_email, docx_file, pdf_file)
                                st.success("🎉 Thank you! Your documents have been emailed to you successfully!")
                                st.balloons()
                            else:
                                st.warning("⚠️ Please enter a valid email address.")

                    else:
                        st.error("❌ Unexpected output format. Try again.")
                except Exception as e:
                    st.error(str(e))
        else:
            st.warning("⚠️ Please paste the Job Description to continue.")
